import customtkinter as ctk
import os
import threading
import json
import docx
from docx.shared import Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class NucleoApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Nucleo Maquetación - Sistema Editorial v3.2 (Affinity)")
        self.geometry("950x700")
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Sidebar
        self.sidebar_frame = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")

        ctk.CTkLabel(self.sidebar_frame, text="NUCLEO v3.2", font=ctk.CTkFont(size=24, weight="bold")).grid(row=0, column=0, padx=20, pady=(30, 10))
        ctk.CTkLabel(self.sidebar_frame, text="Affinity Publisher\nDOCX", font=ctk.CTkFont(size=12)).grid(row=1, column=0, padx=20, pady=10)

        self.nav_btn_extract = ctk.CTkButton(self.sidebar_frame, text="1. EXTRAER", command=lambda: self.select_frame("extract"), fg_color="transparent", border_width=2)
        self.nav_btn_extract.grid(row=2, column=0, padx=20, pady=10, sticky="ew")

        self.nav_btn_generate = ctk.CTkButton(self.sidebar_frame, text="2. GENERAR", command=lambda: self.select_frame("generate"), fg_color="transparent", border_width=2)
        self.nav_btn_generate.grid(row=3, column=0, padx=20, pady=10, sticky="ew")

        self.btn_reset = ctk.CTkButton(self.sidebar_frame, text="RESET", fg_color="#cf3030", command=self.reset_all)
        self.btn_reset.grid(row=5, column=0, padx=20, pady=(40, 10), sticky="ew")

        self.extract_frame = ExtractFrame(self)
        self.generate_frame = GenerateFrame(self)
        self.select_frame("extract")

        # Datos compartidos
        self.extracted_data = None

    def reset_all(self):
        self.extracted_data = None
        self.extract_frame.file_path = None
        self.extract_frame.lbl_file.configure(text="Sin archivo")
        self.extract_frame.btn_run.configure(state="disabled", text="EXTRAER CONTENIDO")
        self.extract_frame.log_box.delete("1.0", "end")
        self.extract_frame.log("Esperando archivo...")
        self.generate_frame.btn_generate.configure(state="disabled", text="GENERAR ARCHIVO")
        self.generate_frame.log_box.delete("1.0", "end")
        self.select_frame("extract")

    def select_frame(self, name):
        self.nav_btn_extract.configure(fg_color="transparent")
        self.nav_btn_generate.configure(fg_color="transparent")
        self.extract_frame.grid_forget()
        self.generate_frame.grid_forget()

        if name == "extract":
            self.extract_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
            self.nav_btn_extract.configure(fg_color=("gray75", "gray25"))
        elif name == "generate":
            self.generate_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
            self.nav_btn_generate.configure(fg_color=("gray75", "gray25"))


class ExtractFrame(ctk.CTkFrame):
    """Frame de extracción de contenido desde DOCX"""
    def __init__(self, master):
        super().__init__(master)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        ctk.CTkLabel(self, text="Paso 1: Extracción a JSON Limpio", font=ctk.CTkFont(size=18, weight="bold")).grid(row=0, column=0, padx=20, pady=20, sticky="w")

        ctrl = ctk.CTkFrame(self)
        ctrl.grid(row=1, column=0, padx=20, pady=10, sticky="ew")

        ctk.CTkButton(ctrl, text="Seleccionar DOCX Original", command=self.select_file).pack(side="left", padx=10, pady=10)
        self.lbl_file = ctk.CTkLabel(ctrl, text="Sin archivo")
        self.lbl_file.pack(side="left", padx=10)

        self.log_box = ctk.CTkTextbox(self)
        self.log_box.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")
        self.log("Esperando archivo...")

        self.btn_run = ctk.CTkButton(self, text="EXTRAER CONTENIDO", state="disabled", fg_color="green", height=40, command=self.run_extract)
        self.btn_run.grid(row=3, column=0, padx=20, pady=20, sticky="ew")

        self.file_path = None

    def log(self, text):
        self.log_box.insert("end", text + "\n")
        self.log_box.see("end")

    def select_file(self):
        path = ctk.filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if path:
            self.file_path = path
            self.lbl_file.configure(text=os.path.basename(path))
            self.btn_run.configure(state="normal")
            self.log(f"Cargado: {path}")

    def run_extract(self):
        self.btn_run.configure(state="disabled")
        threading.Thread(target=self.extract_thread, daemon=True).start()

    def extract_thread(self):
        try:
            self.log("Leyendo documento...")
            doc = docx.Document(self.file_path)

            garbage_patterns = [
                r"^\s*Paso\s+\d",
                r"^\s*Motivos\s+de\s+oraci",
                r"^\s*Testimonios\s+de",
                r"^\s*Lectura\s+del\s+estudio",
                r"^\s*_{5,}",
            ]
            garbage_phrases = [
                "Compartamos en el grupo", "Leamos entre todos", "Si hay personas nuevas",
                "Durante la lectura puedes subrayar", "Compartir mi desafío", "¿Qué frase",
                "¿Que frase", "llamó más tu atención", "Responde honestamente",
                "Recuerda que para crecer", "¿Que área", "¿Qué área",
                "¿Qué acción específica", "Asegúrate de incluir quién",
                "El Espíritu Santo te guiará y te ayudará", "Toma un minuto para compartir",
                "Mientras los compañeros", "Revisamos que hemos apuntado",
                "Es importante también que podamos orar", "Concluimos el estudio",
                "10 minutos", "15 minutos", "30 minutos", "5 minutos",
                "Aplicación", "Desafío semanal"
            ]
            stop_phrases = ["anexo", "apéndice", "guia para el lider", "formato para trabajar"]

            regex_study = re.compile(r"^\s*Estudio\s+(\d+)\s*[-–—:]\s*(.+)", re.IGNORECASE)
            garbage_regex = [re.compile(p, re.IGNORECASE) for p in garbage_patterns]

            def is_garbage(text):
                t = text.strip()
                if not t:
                    return True
                for r in garbage_regex:
                    if r.match(t):
                        return True
                tl = t.lower()
                for ph in garbage_phrases:
                    if ph.lower() in tl:
                        return True
                return False

            def is_stop(text):
                tl = text.lower().strip()
                for sp in stop_phrases:
                    if tl.startswith(sp) or sp in tl:
                        return True
                return False

            studies = []
            current_study = None
            state = "SEARCHING"

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                match = regex_study.match(text)
                if match:
                    if current_study:
                        studies.append(current_study)
                    current_study = {
                        "number": int(match.group(1)),
                        "title": match.group(2).strip(),
                        "content": []
                    }
                    state = "READING"
                    self.log(f"> Estudio {current_study['number']}: {current_study['title'][:30]}...")
                    continue

                if state != "READING":
                    continue

                if is_stop(text):
                    self.log(f"  [STOP] {text[:30]}...")
                    state = "SEARCHING"
                    continue

                if is_garbage(text):
                    self.log(f"  [SKIP] {text[:40]}...")
                    continue

                bold_count = sum(1 for r in para.runs if r.bold and r.text.strip())
                italic_count = sum(1 for r in para.runs if r.italic and r.text.strip())
                total_runs = sum(1 for r in para.runs if r.text.strip())

                if total_runs > 0 and bold_count == total_runs and italic_count == total_runs:
                    content_type = "CITA"
                elif total_runs > 0 and bold_count == total_runs and len(text) < 80:
                    content_type = "SUBTITULO"
                else:
                    content_type = "PARRAFO"

                runs = []
                for r in para.runs:
                    if r.text.strip():
                        runs.append({
                            "text": r.text,
                            "bold": bool(r.bold),
                            "italic": bool(r.italic)
                        })

                current_study["content"].append({
                    "type": content_type,
                    "text": text,
                    "runs": runs
                })

            if current_study:
                studies.append(current_study)

            self.log(f"\n=== EXTRACCIÓN COMPLETA ===")
            self.log(f"Total estudios: {len(studies)}")

            json_path = os.path.splitext(self.file_path)[0] + "_DATOS.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(studies, f, ensure_ascii=False, indent=2)

            self.log(f"Guardado: {json_path}")

            self.master.extracted_data = studies
            self.master.generate_frame.btn_generate.configure(state="normal")

            self.btn_run.configure(state="normal", text="IR AL PASO 2", command=lambda: self.master.select_frame("generate"))

        except Exception as e:
            self.log(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
            self.btn_run.configure(state="normal")


class GenerateFrame(ctk.CTkFrame):
    """Frame de generación de archivos DOCX para Affinity Publisher"""
    def __init__(self, master):
        super().__init__(master)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(4, weight=1)

        ctk.CTkLabel(self, text="Paso 2: Generar Archivos DOCX con Estilos", font=ctk.CTkFont(size=18, weight="bold")).grid(row=0, column=0, columnspan=4, padx=20, pady=(20,10), sticky="w")

        # Estructura del libro
        struct_frame = ctk.CTkFrame(self)
        struct_frame.grid(row=1, column=0, columnspan=4, padx=20, pady=10, sticky="ew")
        ctk.CTkLabel(struct_frame, text="ESTRUCTURA DEL LIBRO:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=10, pady=(10,5))

        self.struct_var = ctk.StringVar(value="single")
        ctk.CTkRadioButton(struct_frame, text="Archivo ÚNICO (todos los estudios juntos) ⭐ RECOMENDADO",
                          variable=self.struct_var, value="single").pack(anchor="w", padx=30, pady=2)
        ctk.CTkRadioButton(struct_frame, text="Archivos SEPARADOS (1 archivo por estudio)",
                          variable=self.struct_var, value="multiple").pack(anchor="w", padx=30, pady=2)

        # Checkbox para saltos de página
        self.page_breaks_var = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(struct_frame, text="Incluir saltos de página entre estudios (desactivar para auto-flow en Affinity)",
                       variable=self.page_breaks_var).pack(anchor="w", padx=30, pady=(5,10))

        # Configuración de página
        style_frame = ctk.CTkFrame(self)
        style_frame.grid(row=2, column=0, columnspan=4, padx=20, pady=10, sticky="ew")
        ctk.CTkLabel(style_frame, text="CONFIGURACIÓN DE PÁGINA:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=10, pady=5, sticky="w")

        self.entries = {}
        self._add_entry(style_frame, "width", "Ancho (mm):", "150", 1, 0)
        self._add_entry(style_frame, "height", "Alto (mm):", "230", 1, 1)

        # Información sobre Affinity
        info_frame = ctk.CTkFrame(self)
        info_frame.grid(row=3, column=0, columnspan=4, padx=20, pady=5, sticky="ew")
        info_text = ("💡 TIP: El archivo DOCX generado es 100% compatible con Affinity V3 (Layout Studio).\n"
                    "Los estilos se importarán automáticamente y podrás editarlos globalmente.")
        ctk.CTkLabel(info_frame, text=info_text, font=ctk.CTkFont(size=11),
                    text_color=("#2563eb", "#3b82f6"), wraplength=800, justify="left").pack(padx=10, pady=10)

        # Log
        self.log_box = ctk.CTkTextbox(self)
        self.log_box.grid(row=4, column=0, columnspan=4, padx=20, pady=10, sticky="nsew")

        # Botón de ayuda para configuración de Affinity
        btn_help_affinity = ctk.CTkButton(self, text="ℹ️ Configuración Recomendada para Affinity",
                                         fg_color="#1e3a8a", hover_color="#1e40af",
                                         command=self.show_affinity_config)
        btn_help_affinity.grid(row=5, column=0, columnspan=4, padx=20, pady=(10,5), sticky="ew")

        # Botón de generación
        self.btn_generate = ctk.CTkButton(self, text="GENERAR ARCHIVO", state="disabled",
                                         fg_color="green", height=40, command=self.generate)
        self.btn_generate.grid(row=6, column=0, columnspan=4, padx=20, pady=(5,20), sticky="ew")

    def _add_entry(self, parent, key, label, default, row, col):
        ctk.CTkLabel(parent, text=label).grid(row=row, column=col*2, padx=10, pady=5, sticky="e")
        e = ctk.CTkEntry(parent, width=80)
        e.insert(0, default)
        e.grid(row=row, column=col*2+1, padx=10, pady=5, sticky="w")
        self.entries[key] = e

    def log(self, text):
        print(text)
        self.after(0, lambda: self._log_safe(text))

    def _log_safe(self, text):
        self.log_box.insert("end", text + "\n")
        self.log_box.see("end")

    def show_affinity_config(self):
        """Muestra ventana con configuración recomendada para Affinity"""
        window = ctk.CTkToplevel(self)
        window.title("Configuración Recomendada - Affinity Publisher")
        window.geometry("800x600")

        # Frame con scroll
        scroll_frame = ctk.CTkScrollableFrame(window, width=760, height=540)
        scroll_frame.pack(padx=20, pady=20, fill="both", expand=True)

        # Título
        ctk.CTkLabel(scroll_frame, text="📐 CONFIGURACIÓN RECOMENDADA PARA AFFINITY PUBLISHER",
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(10,20))

        # Contenido
        config_text = """
Al crear un nuevo documento en Affinity Publisher para importar el DOCX generado:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
AJUSTES DEL DOCUMENTO (Archivo > Nuevo > Ajustes del documento)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ Tipo: Libros CCE

✅ Unidades de documento: Milímetros

✅ PPP (DPI): 300
   (Calidad profesional para impresión)

✅ Anchura de página: 150 mm
   (Coincide con el DOCX generado)

✅ Altura de página: 230 mm
   (Coincide con el DOCX generado)

✅ Formato de color: CMYK/8
   (Para impresión profesional)

✅ Perfil de color: U.S. Web Coated (SWOP) v2
   (Estándar de impresión)

❌ Fondo transparente: Desactivado
   (No necesario para libros impresos)

❌ Mesas de trabajo: Desactivado
   (No necesario para flujo de texto)

✅ Colocación de imágenes: Preferir enlazado
   (Archivos más livianos)

❌ Importar archivos de texto como enlace: Desactivado
   (Queremos texto embebido)

✅ Conservar estilos de texto: ACTIVADO ⭐⭐⭐
   ¡MUY IMPORTANTE! Sin esto los estilos NO se importan

✅ Varias páginas: ACTIVADO
   (Permite crear múltiples páginas automáticamente)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
MÁRGENES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Izquierda: 0 mm  |  Superior: 0 mm
Derecha:   0 mm  |  Inferior: 0 mm

(Los márgenes ya vienen en el DOCX - puedes ajustar después)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SANGRADO (BLEED)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Para libros de texto: 0 mm (todos lados)

Si tu imprenta lo requiere: 3 mm (todos lados)
Consulta primero con tu imprenta.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
IMPORTAR EL DOCX CON AUTO-FLOW
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. Crear documento con la configuración anterior

2. File > Place (Ctrl+Shift+D)

3. Seleccionar: Libro_Completo_XX_Estudios.docx

4. Mantener presionada la tecla SHIFT

5. Hacer UN SOLO CLIC en la primera página

✨ Affinity creará TODAS las páginas automáticamente

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
VERIFICAR ESTILOS IMPORTADOS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Window > Text Styles (Ctrl+Alt+T)

Deberías ver estos 6 estilos:
  • Etiqueta_Estudio
  • Titulo_Principal
  • Cita_Biblica
  • Subtitulo
  • Cuerpo_Texto
  • Cuerpo_Sin_Sangria

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
EDITAR ESTILOS GLOBALMENTE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. Abrir panel: Window > Text Styles

2. Doble clic en cualquier estilo

3. Modificar:
   - Character: fuente, tamaño, color, negrita, cursiva
   - Paragraph: alineación, espaciado, sangrías
   - Advanced: kerning, tracking, etc.

4. Click OK

✨ TODOS los textos con ese estilo cambiarán instantáneamente

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
NOTAS IMPORTANTES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ Si el auto-flow no funciona:
   - Asegúrate de desmarcar el checkbox "Incluir saltos de página"
     en el Paso 2 de esta aplicación
   - Regenera el DOCX
   - Intenta nuevamente en Affinity con SHIFT + clic

⚠️ Si los estilos no se importan:
   - Verifica que "Conservar estilos de texto" esté ACTIVADO
   - Usa File > Place (NO File > Open)

⚠️ Fuente "Roble" no encontrada:
   - Opción A: Instalar fuente Roble.ttf
   - Opción B: Editar estilos en Affinity y cambiar a Arial

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💡 TIP: Guarda tu documento de Affinity como plantilla (.aftemplate)
   para reutilizar la configuración en futuros libros.
"""

        text_widget = ctk.CTkTextbox(scroll_frame, width=740, height=450,
                                     font=ctk.CTkFont(family="Consolas", size=11))
        text_widget.pack(pady=10, padx=10)
        text_widget.insert("1.0", config_text)
        text_widget.configure(state="disabled")

        # Botón cerrar
        ctk.CTkButton(window, text="Cerrar", command=window.destroy,
                     fg_color="#dc2626").pack(pady=10)

    def generate(self):
        self.btn_generate.configure(state="disabled")
        threading.Thread(target=self._gen_thread, daemon=True).start()

    def _gen_thread(self):
        try:
            studies = self.master.extracted_data
            if not studies:
                self.log("ERROR: No hay datos. Ejecute el Paso 1 primero.")
                self.btn_generate.configure(state="normal")
                return

            structure = self.struct_var.get()

            output_dir = os.path.join(os.getcwd(), "Salida_Affinity_v3")
            os.makedirs(output_dir, exist_ok=True)

            if structure == "single":
                self._generate_single_docx(studies, output_dir)
            else:
                self._generate_multiple_docx(studies, output_dir)

            self.log(f"\n✅ GENERACIÓN COMPLETADA")
            self.btn_generate.configure(state="normal", text="ABRIR CARPETA",
                                       command=lambda: os.startfile(output_dir))

        except Exception as e:
            self.log(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
            self.btn_generate.configure(state="normal")

    def _generate_single_docx(self, studies, output_dir):
        """Genera UN SOLO archivo DOCX con todos los estudios"""
        self.log(f"Generando archivo DOCX único con {len(studies)} estudios...")

        doc = docx.Document()

        # Configurar página
        section = doc.sections[0]
        section.page_width = Mm(float(self.entries["width"].get()))
        section.page_height = Mm(float(self.entries["height"].get()))
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

        # Crear estilos
        self._create_docx_styles(doc)

        # Agregar cada estudio
        for idx, study in enumerate(studies):
            if idx > 0 and self.page_breaks_var.get():
                doc.add_page_break()  # Salto de página entre estudios (solo si checkbox activado)

            # Etiqueta
            p = doc.add_paragraph(f"ESTUDIO {study['number']}", style="Etiqueta_Estudio")

            # Título
            p = doc.add_paragraph(study['title'], style="Titulo_Principal")

            # Contenido
            last_type = "TITULO"
            for item in study["content"]:
                if item["type"] == "CITA":
                    p = doc.add_paragraph(item['text'], style="Cita_Biblica")
                    last_type = "CITA"
                elif item["type"] == "SUBTITULO":
                    p = doc.add_paragraph(item['text'], style="Subtitulo")
                    last_type = "SUBTITULO"
                else:
                    style = "Cuerpo_Sin_Sangria" if last_type in ["TITULO", "SUBTITULO", "CITA"] else "Cuerpo_Texto"
                    p = doc.add_paragraph()
                    p.style = doc.styles[style]
                    for run_data in item["runs"]:
                        r = p.add_run(run_data['text'])
                        r.bold = run_data['bold']
                        r.italic = run_data['italic']
                    last_type = "PARRAFO"

            self.log(f"  ✓ Estudio {study['number']}: {study['title'][:40]}...")

        # Guardar
        fname = f"Libro_Completo_{len(studies)}_Estudios.docx"
        fpath = os.path.join(output_dir, fname)
        doc.save(fpath)

        self.log(f"\n📄 Archivo guardado: {fname}")
        self.log(f"📊 Total: {len(studies)} estudios en 1 archivo")
        self.log(f"\n🎯 PARA AFFINITY V3:")
        self.log(f"1. Abrir Affinity (Layout Studio)")
        self.log(f"2. File > Place > Seleccionar '{fname}'")
        self.log(f"3. Los estilos se importarán automáticamente")
        self.log(f"4. Editar estilos: Window > Text Styles")

    def _generate_multiple_docx(self, studies, output_dir):
        """Genera archivos DOCX separados (1 por estudio)"""
        self.log(f"Generando {len(studies)} archivos DOCX separados...")

        for study in studies:
            doc = docx.Document()

            section = doc.sections[0]
            section.page_width = Mm(float(self.entries["width"].get()))
            section.page_height = Mm(float(self.entries["height"].get()))
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)

            self._create_docx_styles(doc)

            p = doc.add_paragraph(f"ESTUDIO {study['number']}", style="Etiqueta_Estudio")
            p = doc.add_paragraph(study['title'], style="Titulo_Principal")

            last_type = "TITULO"
            for item in study["content"]:
                if item["type"] == "CITA":
                    p = doc.add_paragraph(item['text'], style="Cita_Biblica")
                    last_type = "CITA"
                elif item["type"] == "SUBTITULO":
                    p = doc.add_paragraph(item['text'], style="Subtitulo")
                    last_type = "SUBTITULO"
                else:
                    style = "Cuerpo_Sin_Sangria" if last_type in ["TITULO", "SUBTITULO", "CITA"] else "Cuerpo_Texto"
                    p = doc.add_paragraph()
                    p.style = doc.styles[style]
                    for run_data in item["runs"]:
                        r = p.add_run(run_data['text'])
                        r.bold = run_data['bold']
                        r.italic = run_data['italic']
                    last_type = "PARRAFO"

            safe_title = "".join(c if c.isalnum() or c in " _" else "" for c in study["title"]).strip().replace(" ", "_")[:30]
            fname = f"{study['number']:02d}_{safe_title}.docx"
            fpath = os.path.join(output_dir, fname)
            doc.save(fpath)

            self.log(f"  ✓ {fname}")

        self.log(f"\n📊 Total: {len(studies)} archivos generados")

    def _create_docx_styles(self, doc):
        """Crea estilos nativos de Word compatibles con Affinity"""
        styles = doc.styles

        # Estilo Etiqueta_Estudio
        try:
            style = styles.add_style('Etiqueta_Estudio', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Etiqueta_Estudio']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.font.color.rgb = docx.shared.RGBColor(102, 102, 102)  # #666666
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.space_before = Pt(0)

        # Estilo Titulo_Principal
        try:
            style = styles.add_style('Titulo_Principal', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Titulo_Principal']
        style.font.name = 'Arial'
        style.font.size = Pt(20)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(184)

        # Estilo Cita_Biblica
        try:
            style = styles.add_style('Cita_Biblica', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Cita_Biblica']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.left_indent = Mm(10)
        style.paragraph_format.right_indent = Mm(10)
        style.paragraph_format.space_after = Pt(17)

        # Estilo Referencia_Biblica (referencias de versículos)
        try:
            style = styles.add_style('Referencia_Biblica', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Referencia_Biblica']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)

        # Estilo Subtitulo
        try:
            style = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Subtitulo']
        style.font.name = 'Arial'
        style.font.size = Pt(13)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(17)
        style.paragraph_format.space_after = Pt(6)

        # Estilo Cuerpo_Texto
        try:
            style = styles.add_style('Cuerpo_Texto', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Cuerpo_Texto']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.first_line_indent = Mm(5)
        style.paragraph_format.space_after = Pt(11)

        # Estilo Cuerpo_Sin_Sangria
        try:
            style = styles.add_style('Cuerpo_Sin_Sangria', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Cuerpo_Sin_Sangria']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(11)

        # Estilo Paso_Estudio (con viñeta)
        try:
            style = styles.add_style('Paso_Estudio', WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles['Paso_Estudio']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Mm(5)


if __name__ == "__main__":
    app = NucleoApp()
    app.mainloop()
