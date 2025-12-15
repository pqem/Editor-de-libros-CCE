# -*- coding: utf-8 -*-
"""Script de prueba para verificar que los nuevos estilos se crean correctamente"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_estilos_prueba():
    """Crea un documento de prueba con todos los estilos"""
    doc = Document()
    styles = doc.styles

    # Lista de estilos a crear (copiado de Nucleo_App_v3_Affinity.py)

    # 1. Etiqueta_Estudio
    try:
        style = styles.add_style('Etiqueta_Estudio', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Etiqueta_Estudio']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(102, 102, 102)  # #666666
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.space_before = Pt(0)

    # 2. Titulo_Principal
    try:
        style = styles.add_style('Titulo_Principal', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Titulo_Principal']
    style.font.name = 'Arial'
    style.font.size = Pt(20)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(184)

    # 3. Cita_Biblica
    try:
        style = styles.add_style('Cita_Biblica', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Cita_Biblica']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.italic = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.left_indent = Mm(10)
    style.paragraph_format.right_indent = Mm(10)
    style.paragraph_format.space_after = Pt(17)

    # 4. Referencia_Biblica (NUEVO)
    try:
        style = styles.add_style('Referencia_Biblica', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Referencia_Biblica']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.font.italic = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    # 5. Subtitulo
    try:
        style = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Subtitulo']
    style.font.name = 'Arial'
    style.font.size = Pt(13)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(17)
    style.paragraph_format.space_after = Pt(6)

    # 6. Cuerpo_Texto
    try:
        style = styles.add_style('Cuerpo_Texto', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Cuerpo_Texto']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Mm(5)
    style.paragraph_format.space_after = Pt(11)

    # 7. Cuerpo_Sin_Sangria
    try:
        style = styles.add_style('Cuerpo_Sin_Sangria', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Cuerpo_Sin_Sangria']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(11)

    # 8. Paso_Estudio (NUEVO)
    try:
        style = styles.add_style('Paso_Estudio', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles['Paso_Estudio']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.left_indent = Mm(5)

    # Añadir contenido de ejemplo
    doc.add_paragraph("ESTUDIO 1", style="Etiqueta_Estudio")
    doc.add_paragraph("El estallido del Reino", style="Titulo_Principal")
    doc.add_paragraph("Cuando llegó el día de Pentecostés, estaban todos juntos en el mismo lugar...", style="Cita_Biblica")
    doc.add_paragraph("(Hechos 2:1-4, NVI)", style="Referencia_Biblica")
    doc.add_paragraph("La unidad precede a la manifestación", style="Subtitulo")
    doc.add_paragraph("Este es el primer párrafo sin sangría después del subtítulo. Contiene información importante.", style="Cuerpo_Sin_Sangria")
    doc.add_paragraph("Este es el segundo párrafo con sangría. Continúa el desarrollo del tema anterior.", style="Cuerpo_Texto")
    doc.add_paragraph("• Paso 1 (10 min) - Testimonios de autoevaluación", style="Paso_Estudio")
    doc.add_paragraph("Instrucciones para el paso 1 del estudio.", style="Cuerpo_Sin_Sangria")

    # Guardar
    output_path = "TEST_Estilos_Nuevos.docx"
    doc.save(output_path)
    print(f"✅ Documento de prueba creado: {output_path}")

    # Verificar estilos creados
    print("\n📋 Estilos verificados:")
    for style_name in ['Etiqueta_Estudio', 'Titulo_Principal', 'Cita_Biblica', 'Referencia_Biblica',
                       'Subtitulo', 'Cuerpo_Texto', 'Cuerpo_Sin_Sangria', 'Paso_Estudio']:
        if style_name in [s.name for s in doc.styles]:
            print(f"  ✓ {style_name}")
        else:
            print(f"  ✗ {style_name} - NO ENCONTRADO")

if __name__ == "__main__":
    crear_estilos_prueba()
