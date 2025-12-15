# -*- coding: utf-8 -*-
from docx import Document
import glob

# Find DOCX file
docx_files = glob.glob("*.docx")
print(f"Archivos DOCX encontrados: {docx_files}")

if docx_files:
    doc = Document(docx_files[0])
    styles = set()
    content_samples = {}

    print(f"\nAnalizando: {docx_files[0]}")
    print(f"Total de párrafos: {len(doc.paragraphs)}\n")

    for i, para in enumerate(doc.paragraphs[:200]):
        text = para.text.strip()
        if text:
            style = para.style.name
            styles.add(style)
            if style not in content_samples:
                content_samples[style] = text[:150]

    print("=" * 60)
    print("ESTILOS ENCONTRADOS EN EL DOCUMENTO")
    print("=" * 60)
    for s in sorted(styles):
        print(f"\n📌 Estilo: {s}")
        if s in content_samples:
            print(f"   Ejemplo: {content_samples[s]}")

    print(f"\n\n✅ Total de estilos únicos: {len(styles)}")

    # Buscar secciones especiales
    print("\n" + "=" * 60)
    print("SECCIONES ESPECIALES DETECTADAS")
    print("=" * 60)

    special_sections = {
        'Prólogo': [],
        'Introducción': [],
        'Referencias': [],
        'Notas': [],
        'Conclusión': []
    }

    for para in doc.paragraphs[:200]:
        text = para.text.strip().lower()
        for section_name in special_sections.keys():
            if section_name.lower() in text:
                special_sections[section_name].append(para.text[:100])

    for section, examples in special_sections.items():
        if examples:
            print(f"\n🔖 {section}: {len(examples)} ocurrencia(s)")
            print(f"   {examples[0]}")
